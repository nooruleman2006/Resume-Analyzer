"""
Resume Parser Service
Handles text extraction from PDF, DOCX, and legacy DOC files.
"""

import os
# REPLACE with this:
from pypdf import PdfReader

def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return '\n'.join(pages).strip()

# ── DOCX ──────────────────────────────────────────────────────
def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file using python-docx.
    Captures body paragraphs, tables, and section headers/footers.
    Falls back to raw ZIP/XML parsing if python-docx is unavailable.
    """
    try:
        from docx import Document

        doc   = Document(file_path)
        parts = []

        def _para_text(para):
            return ''.join(run.text for run in para.runs).strip()

        # Headers / Footers
        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf and not hf.is_linked_to_previous:
                    for para in hf.paragraphs:
                        t = _para_text(para)
                        if t:
                            parts.append(t)

        # Body paragraphs + tables (in document order)
        for block in doc.element.body:
            tag = block.tag.split('}')[-1]

            if tag == 'p':
                from docx.text.paragraph import Paragraph
                para = Paragraph(block, doc)
                t = _para_text(para)
                if t:
                    parts.append(t)

            elif tag == 'tbl':
                from docx.table import Table
                tbl = Table(block, doc)
                for row in tbl.rows:
                    row_parts = []
                    for cell in row.cells:
                        cell_t = ' '.join(
                            _para_text(p) for p in cell.paragraphs if _para_text(p)
                        )
                        if cell_t:
                            row_parts.append(cell_t)
                    if row_parts:
                        parts.append(' | '.join(row_parts))

        return '\n'.join(parts).strip()

    except ImportError:
        return _extract_docx_xml_fallback(file_path)


def _extract_docx_xml_fallback(file_path):
    """Fallback DOCX extractor using only stdlib — no python-docx needed."""
    import zipfile
    import xml.etree.ElementTree as ET

    NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    parts = []

    with zipfile.ZipFile(file_path, 'r') as z:
        xml_files = [
            name for name in z.namelist()
            if name.startswith('word/') and name.endswith('.xml')
            and any(s in name for s in ('document', 'header', 'footer'))
        ]
        for xml_file in sorted(xml_files):
            with z.open(xml_file) as f:
                root = ET.parse(f).getroot()
                for para in root.iter(f'{{{NS}}}p'):
                    runs = para.findall(f'.//{{{NS}}}t')
                    line = ' '.join(r.text for r in runs if r.text)
                    if line.strip():
                        parts.append(line.strip())

    return '\n'.join(parts)


# ── DOC (legacy OLE2) ─────────────────────────────────────────
def extract_text_from_doc(file_path):
    """
    Extract text from a legacy .doc file.
    Tries antiword → docx2txt → raw binary scan.
    """
    # 1. antiword (best quality if available)
    try:
        import subprocess
        result = subprocess.run(
            ['antiword', file_path],
            capture_output=True, text=True, timeout=15
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except (FileNotFoundError, Exception):
        pass

    # 2. docx2txt
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        if text and text.strip():
            return text.strip()
    except Exception:
        pass

    # 3. Raw binary ASCII scan (last resort)
    try:
        import re
        with open(file_path, 'rb') as f:
            raw = f.read()
        strings = re.findall(rb'[ -~]{4,}', raw)
        return '\n'.join(s.decode('ascii', errors='ignore') for s in strings).strip()
    except Exception:
        return ''


# ── Public API ────────────────────────────────────────────────
def parse_resume(file_path):
    """
    Auto-detect file type and extract text. Returns raw text string.
    Raises ValueError for unsupported file types.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.doc':
        return extract_text_from_doc(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
