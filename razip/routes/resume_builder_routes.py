from flask import (Blueprint, render_template, request, redirect, url_for,
                   flash, session, jsonify, send_file, make_response)
from flask_login import login_required, current_user
from models.db import query
from services.ats_scorer import calculate_ats_score
from services.ai_analyzer import analyze_resume
from models.resume import Resume
import json, io, os, re
from datetime import datetime

builder_bp = Blueprint('builder', __name__)

# ── Template catalogue ─────────────────────────────────────
TEMPLATES_LIST = [
    {'id': 'professional',    'label': 'Professional',      'file': 'resume_professional.html',    'preview_color': '#7c3aed', 'description': 'Clean corporate style'},
    {'id': 'executive',       'label': 'Executive',         'file': 'resume_executive.html',       'preview_color': '#0f4c3a', 'description': 'Bold executive presence'},
    {'id': 'modern_minimal',  'label': 'Modern Minimal',    'file': 'resume_modern_minimal.html',  'preview_color': '#2563eb', 'description': 'Sleek and contemporary'},
    {'id': 'simple_classic',  'label': 'Simple Classic',    'file': 'resume_simple_classic.html',  'preview_color': '#374151', 'description': 'Timeless classic layout'},
    {'id': 'elegant',         'label': 'Elegant',           'file': 'resume_elegant.html',         'preview_color': '#9f7aea', 'description': 'Refined elegant design'},
    {'id': 'tech_developer',  'label': 'Tech Developer',    'file': 'resume_tech_developer.html',  'preview_color': '#0ea5e9', 'description': 'Perfect for engineers'},
    {'id': 'creative_bold',   'label': 'Creative Bold',     'file': 'resume_creative_bold.html',   'preview_color': '#e11d48', 'description': 'Stand-out creative flair'},
    {'id': 'academic',        'label': 'Academic',          'file': 'resume_academic.html',        'preview_color': '#1d4ed8', 'description': 'Ideal for researchers'},
    {'id': 'ats_friendly',    'label': 'ATS Friendly',      'file': 'resume_ats_friendly.html',    'preview_color': '#15803d', 'description': 'Optimised for ATS systems'},
    {'id': 'two_column_pro',  'label': 'Two Column Pro',    'file': 'resume_two_column_pro.html',  'preview_color': '#0f766e', 'description': 'Space-efficient two columns'},
    {'id': 'fresh_graduate',  'label': 'Fresh Graduate',    'file': 'resume_fresh_graduate.html',  'preview_color': '#d97706', 'description': 'Great for new graduates'},
    {'id': 'corporate_blue',  'label': 'Corporate Blue',    'file': 'resume_corporate_blue.html',  'preview_color': '#1e40af', 'description': 'Traditional corporate blue'},
    {'id': 'dark_mode',       'label': 'Dark Mode',         'file': 'resume_dark_mode.html',       'preview_color': '#1f2937', 'description': 'Striking dark theme'},
    {'id': 'compact_one_page','label': 'Compact One Page',  'file': 'resume_compact_one_page.html','preview_color': '#6d28d9', 'description': 'Fits everything on one page'},
    {'id': 'infographic',     'label': 'Infographic',       'file': 'resume_infographic.html',     'preview_color': '#db2777', 'description': 'Visual infographic style'},
    {'id': 'minimalist_pro',  'label': 'Minimalist Pro',    'file': 'resume_minimalist_pro.html',  'preview_color': '#475569', 'description': 'Clean minimalist layout'},
    {'id': 'slate_gold',      'label': 'Slate & Gold',      'file': 'resume_slate_gold.html',      'preview_color': '#b45309', 'description': 'Sophisticated slate & gold'},
]

_TEMPLATES_BY_ID = {t['id']: t for t in TEMPLATES_LIST}


def _get_template(template_id):
    """Return template dict, falling back to 'professional'."""
    return _TEMPLATES_BY_ID.get(template_id) or _TEMPLATES_BY_ID['professional']


def _get_draft(draft_id):
    return query(
        "SELECT * FROM resume_builder_drafts WHERE id=? AND user_id=?",
        (draft_id, current_user.id), fetchone=True
    )


# ── List / landing ────────────────────────────────────────────
@builder_bp.route('/resume-builder')
@login_required
def builder_home():
    drafts = query(
        "SELECT * FROM resume_builder_drafts WHERE user_id=? ORDER BY updated_at DESC",
        (current_user.id,), fetchall=True
    )
    return render_template('resume_builder_list.html', drafts=drafts)


# ── New draft ─────────────────────────────────────────────────
@builder_bp.route('/resume-builder/new', methods=['POST'])
@login_required
def builder_new():
    name = request.form.get('name', 'My Resume').strip() or 'My Resume'
    draft_id = query(
        "INSERT INTO resume_builder_drafts (user_id, name, data, template) VALUES (?,?,?,?)",
        (current_user.id, name, '{}', 'professional'), commit=True
    )
    return redirect(url_for('builder.builder_edit', draft_id=draft_id))


# ── Edit draft ────────────────────────────────────────────────
@builder_bp.route('/resume-builder/<int:draft_id>/edit')
@login_required
def builder_edit(draft_id):
    draft = _get_draft(draft_id)
    if not draft:
        flash('Resume draft not found.', 'error')
        return redirect(url_for('builder.builder_home'))
    try:
        data = json.loads(draft['data'] or '{}')
    except Exception:
        data = {}
    return render_template('resume_builder.html', draft=draft, data=data,
                           templates=TEMPLATES_LIST)


# ── HTML Preview (renders selected Jinja2 template) ───────────
@builder_bp.route('/resume-builder/<int:draft_id>/preview')
@login_required
def builder_preview(draft_id):
    draft = _get_draft(draft_id)
    if not draft:
        flash('Draft not found.', 'error')
        return redirect(url_for('builder.builder_home'))

    try:
        data = json.loads(draft['data'] or '{}')
    except Exception:
        data = {}

    template_id = request.args.get('template_id') or draft.get('template') or 'professional'
    tpl = _get_template(template_id)

    personal       = data.get('personal', {})
    summary        = data.get('summary', '')
    experience     = data.get('experience', [])
    education      = data.get('education', [])
    skills         = data.get('skills', [])
    certifications = data.get('certifications', [])
    languages      = data.get('languages', [])
    projects       = data.get('projects', [])

    return render_template(
        f'resume_builder/templates/{tpl["file"]}',
        personal=personal, summary=summary, experience=experience,
        education=education, skills=skills, certifications=certifications,
        languages=languages, projects=projects
    )


# ── AJAX save ─────────────────────────────────────────────────
@builder_bp.route('/resume-builder/<int:draft_id>/save', methods=['POST'])
@login_required
def builder_save(draft_id):
    draft = _get_draft(draft_id)
    if not draft:
        return jsonify({'ok': False, 'error': 'Not found'}), 404

    payload = request.get_json(silent=True) or {}
    data     = json.dumps(payload.get('data', {}))
    template = payload.get('template', 'professional')
    name     = payload.get('name', draft['name'])[:80]

    query(
        "UPDATE resume_builder_drafts SET data=?, template=?, name=?, updated_at=datetime('now') WHERE id=?",
        (data, template, name, draft_id), commit=True
    )
    return jsonify({'ok': True})


# ── Delete draft ──────────────────────────────────────────────
@builder_bp.route('/resume-builder/<int:draft_id>/delete', methods=['POST'])
@login_required
def builder_delete(draft_id):
    draft = _get_draft(draft_id)
    if draft:
        query("DELETE FROM resume_builder_drafts WHERE id=?", (draft_id,), commit=True)
        flash('Resume deleted.', 'success')
    return redirect(url_for('builder.builder_home'))


# ── Analyze built resume ───────────────────────────────────────
@builder_bp.route('/resume-builder/<int:draft_id>/analyze', methods=['POST'])
@login_required
def builder_analyze(draft_id):
    draft = _get_draft(draft_id)
    if not draft:
        flash('Draft not found.', 'error')
        return redirect(url_for('builder.builder_home'))

    try:
        data = json.loads(draft['data'] or '{}')
    except Exception:
        data = {}

    resume_text = _data_to_text(data)
    if len(resume_text.strip()) < 30:
        flash('Please fill in your resume before analyzing.', 'error')
        return redirect(url_for('builder.builder_edit', draft_id=draft_id))

    ats_data = calculate_ats_score(resume_text)
    ai_data  = analyze_resume(resume_text,
                              present_skills=ats_data['present_skills'],
                              missing_skills=ats_data['missing_skills'])

    name = data.get('personal', {}).get('name', 'Built Resume') or 'Built Resume'
    filename = f"{name.replace(' ', '_')}_Builder.txt"

    resume_id = Resume.create(
        user_id         = current_user.id,
        filename        = filename,
        stored_path     = f'builder_{draft_id}',
        raw_text        = resume_text,
        skills          = ats_data['present_skills'],
        job_title       = '',
        industry        = '',
        job_description = '',
    )
    query(
        """INSERT INTO analysis_results
           (resume_id, ats_score, readability_score, keyword_score,
            ai_summary, ai_suggestions, ai_strengths, ai_weaknesses,
            ai_suitable_roles, present_skills, missing_skills, score_breakdown)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            resume_id,
            ats_data['ats_score'], ats_data['readability_score'], ats_data['keyword_score'],
            ai_data.get('summary', ''),
            json.dumps(ai_data.get('suggestions', [])),
            json.dumps(ai_data.get('strengths', [])),
            json.dumps(ai_data.get('weaknesses', [])),
            json.dumps(ai_data.get('suitable_roles', [])),
            json.dumps(ats_data['present_skills']),
            json.dumps(ats_data['missing_skills']),
            json.dumps(ats_data['score_breakdown']),
        ),
        commit=True
    )
    session['last_resume_id'] = resume_id
    flash('Resume analyzed! View your full ATS report below.', 'success')
    return redirect(url_for('resume.analysis'))


# ── Export PDF ────────────────────────────────────────────────
@builder_bp.route('/resume-builder/<int:draft_id>/export/pdf')
@login_required
def builder_export_pdf(draft_id):
    draft = _get_draft(draft_id)
    if not draft:
        flash('Draft not found.', 'error')
        return redirect(url_for('builder.builder_home'))

    try:
        data = json.loads(draft['data'] or '{}')
    except Exception:
        data = {}

    template = draft.get('template') or data.get('template') or 'professional'
    # sanitize: only allow known template ids
    if template not in _TEMPLATES_BY_ID:
        template = 'professional'

    pdf_bytes = _generate_pdf(data, template)
    name      = data.get('personal', {}).get('name', 'Resume').replace(' ', '_')
    filename  = f"{name}_Resume.pdf"

    response = make_response(pdf_bytes)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ── Export DOCX ───────────────────────────────────────────────
@builder_bp.route('/resume-builder/<int:draft_id>/export/docx')
@login_required
def builder_export_docx(draft_id):
    draft = _get_draft(draft_id)
    if not draft:
        flash('Draft not found.', 'error')
        return redirect(url_for('builder.builder_home'))

    try:
        data = json.loads(draft['data'] or '{}')
    except Exception:
        data = {}

    docx_bytes = _generate_docx(data)
    name       = data.get('personal', {}).get('name', 'Resume').replace(' ', '_')
    filename   = f"{name}_Resume.docx"

    return send_file(
        io.BytesIO(docx_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


# ── Helpers ───────────────────────────────────────────────────

def _data_to_text(data):
    """Convert builder data dict to plain resume text for ATS scoring."""
    lines = []
    p = data.get('personal', {})
    if p.get('name'):    lines.append(p['name'])
    if p.get('email'):   lines.append(p['email'])
    if p.get('phone'):   lines.append(p['phone'])
    if p.get('location'):lines.append(p['location'])
    if p.get('linkedin'):lines.append(p['linkedin'])

    if data.get('summary'):
        lines += ['', 'SUMMARY', data['summary']]

    if data.get('experience'):
        lines.append('')
        lines.append('EXPERIENCE')
        for e in data['experience']:
            lines.append(f"{e.get('role','')} at {e.get('company','')}")
            lines.append(f"{e.get('start','')} – {e.get('end','Present')}")
            if e.get('description'):
                lines.append(e['description'])

    if data.get('education'):
        lines.append('')
        lines.append('EDUCATION')
        for e in data['education']:
            lines.append(f"{e.get('degree','')} in {e.get('field','')} – {e.get('institution','')}")
            if e.get('year'): lines.append(str(e['year']))

    if data.get('skills'):
        lines += ['', 'SKILLS', ', '.join(data['skills'])]

    if data.get('projects'):
        lines.append('')
        lines.append('PROJECTS')
        for p in data['projects']:
            lines.append(p.get('name', ''))
            if p.get('description'): lines.append(p['description'])
            if p.get('technologies'): lines.append(f"Technologies: {p['technologies']}")

    if data.get('certifications'):
        lines.append('')
        lines.append('CERTIFICATIONS')
        for c in data['certifications']:
            lines.append(f"{c.get('name','')} — {c.get('org','')} {c.get('year','')}")

    if data.get('languages'):
        lines.append('')
        lines.append('LANGUAGES')
        for l in data['languages']:
            lines.append(f"{l.get('language','')} ({l.get('level','')})")

    return '\n'.join(lines)


def _generate_pdf(data, template='professional'):
    """Generate a styled PDF using fpdf2 with per-template designs."""
    from fpdf import FPDF
    import io

    def h2r(hex_color):
        h = hex_color.lstrip('#')
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    class ResumePDF(FPDF):
        pass

    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    PW, M = 210, 15          # page width, margin
    CW = PW - 2 * M          # content width = 180
    LEFT_W, GAP, RIGHT_W = 60, 8, 112
    LEFT_X, RIGHT_X = M, M + LEFT_W + GAP

    p = data.get('personal', {})
    name = p.get('name', '')
    contacts = [c for c in [p.get('email'), p.get('phone'), p.get('location'),
                            p.get('linkedin'), p.get('portfolio')] if c]

    def sec_title(text, color, y, w, size=10, line_color=None):
        pdf.set_xy(M, y)
        pdf.set_font('Helvetica', 'B', size)
        pdf.set_text_color(*h2r(color))
        pdf.cell(w, 6, text, ln=1)
        pdf.set_draw_color(*h2r(line_color or color))
        pdf.set_line_width(0.4)
        pdf.line(M, y + 6, M + w, y + 6)
        return y + 9

    def entry_block(title, subtitle, desc, y, w, title_bold=True, title_size=9.5,
                    title_color='#1e1b4b', sub_color='#6b7280', body_color='#374151'):
        pdf.set_xy(M, y)
        pdf.set_font('Helvetica', 'B' if title_bold else '', title_size)
        pdf.set_text_color(*h2r(title_color))
        pdf.cell(w, 5, title, ln=1)
        if subtitle:
            pdf.set_font('Helvetica', 'I', 8)
            pdf.set_text_color(*h2r(sub_color))
            pdf.cell(w, 4, subtitle, ln=1)
        if desc:
            pdf.set_font('Helvetica', '', 8.5)
            pdf.set_text_color(*h2r(body_color))
            for line in desc.split('\n'):
                line = line.strip().lstrip('* ').lstrip('-').strip()
                if line:
                    pdf.multi_cell(w, 4, f"*  {line}")
        return pdf.get_y() + 2

    def chip(text, bg, fg, y, x=None):
        x = x or pdf.get_x()
        pdf.set_xy(x, y)
        pdf.set_fill_color(*h2r(bg))
        pdf.set_text_color(*h2r(fg))
        pdf.set_font('Helvetica', 'B', 7)
        tw = pdf.get_string_width(text) + 4
        pdf.cell(tw, 5, f'  {text}', ln=0, fill=True)
        pdf.set_xy(x + tw + 2, y)
        return x + tw + 2, y

    # ═══════════════════════════════════════════════════════════
    # COLORFUL PRO
    # ═══════════════════════════════════════════════════════════
    if template == 'colorful_pro':
        pdf.set_fill_color(*h2r('#1a1a2e'))
        pdf.rect(M, 10, CW, 22, style='F')
        pdf.set_xy(M + 4, 14)
        pdf.set_font('Helvetica', 'B', 16)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(CW - 8, 7, name, ln=1)
        pdf.set_font('Helvetica', '', 7)
        pdf.set_text_color(*h2r('#c7d2fe'))
        pdf.set_xy(M + 4, 22)
        pdf.cell(CW - 8, 4, '  |  '.join(contacts[:4]), ln=1)
        # ATS badge
        pdf.set_fill_color(*h2r('#14b8a6'))
        pdf.set_text_color(255, 255, 255)
        pdf.set_font('Helvetica', 'B', 7)
        pdf.set_xy(PW - M - 30, 13)
        pdf.cell(26, 5, ' ATS 92%', ln=1, fill=True)

        # Sidebar background
        pdf.set_fill_color(*h2r('#0f0f23'))
        pdf.rect(M, 36, LEFT_W, 250, style='F')
        ly, ry = 40, 40

        # ── Sidebar ──
        if data.get('skills'):
            pdf.set_xy(LEFT_X + 3, ly)
            pdf.set_font('Helvetica', 'B', 7)
            pdf.set_text_color(*h2r('#14b8a6'))
            pdf.cell(LEFT_W - 6, 5, 'SKILLS', ln=1)
            ly += 7
            sk_colors = [('#818cf8', '#1e1b4b'), ('#14b8a6', '#0f0f23'),
                         ('#f59e0b', '#1e1b4b'), ('#ec4899', '#1e1b4b')]
            for i, s in enumerate(data['skills']):
                fg, bg = sk_colors[i % 4]
                pdf.set_xy(LEFT_X + 3, ly)
                pdf.set_fill_color(*h2r(bg))
                pdf.set_text_color(*h2r(fg))
                pdf.set_font('Helvetica', 'B', 7)
                tw = pdf.get_string_width(s) + 5
                pdf.cell(tw, 4, f'  {s}  ', ln=1, fill=True)
                ly += 5
            ly += 2

        if data.get('education'):
            pdf.set_xy(LEFT_X + 3, ly)
            pdf.set_font('Helvetica', 'B', 7)
            pdf.set_text_color(*h2r('#14b8a6'))
            pdf.cell(LEFT_W - 6, 5, 'EDUCATION', ln=1)
            ly += 7
            for e in data['education']:
                deg = f"{e.get('degree','')} in {e.get('field','')}" if e.get('field') else e.get('degree','')
                pdf.set_xy(LEFT_X + 3, ly)
                pdf.set_font('Helvetica', 'B', 8)
                pdf.set_text_color(*h2r('#e5e7eb'))
                pdf.cell(LEFT_W - 6, 4, e.get('institution',''), ln=1)
                ly += 4
                pdf.set_xy(LEFT_X + 3, ly)
                pdf.set_font('Helvetica', '', 7)
                pdf.set_text_color(*h2r('#9ca3af'))
                yr = f' {e["year"]}' if e.get('year') else ''
                pdf.cell(LEFT_W - 6, 4, f"{deg}{yr}", ln=1)
                ly += 6
            ly += 2

        if data.get('languages'):
            pdf.set_xy(LEFT_X + 3, ly)
            pdf.set_font('Helvetica', 'B', 7)
            pdf.set_text_color(*h2r('#14b8a6'))
            pdf.cell(LEFT_W - 6, 5, 'LANGUAGES', ln=1)
            ly += 7
            levels = {'Native':100,'Fluent':90,'Advanced':75,'Intermediate':55,'Basic':35}
            for l in data['languages']:
                pdf.set_xy(LEFT_X + 3, ly)
                pdf.set_font('Helvetica', '', 7.5)
                pdf.set_text_color(*h2r('#d1d5db'))
                pdf.cell(LEFT_W - 6, 4, f"{l.get('language','')}  {l.get('level','')}", ln=1)
                ly += 4
                bar_w = LEFT_W - 10
                fill_w = bar_w * levels.get(l.get('level',''), 60) / 100
                pdf.set_fill_color(*h2r('#1a1a2e'))
                pdf.rect(LEFT_X + 3, ly, bar_w, 2, style='F')
                pdf.set_fill_color(*h2r('#6366f1'))
                pdf.rect(LEFT_X + 3, ly, fill_w, 2, style='F')
                ly += 4
            ly += 2

        # ── Main ──
        if data.get('summary'):
            ry = sec_title('SUMMARY', '#6366f1', ry, RIGHT_W, size=9, line_color='#6366f1')
            pdf.set_xy(RIGHT_X, ry)
            pdf.set_font('Helvetica', '', 8.5)
            pdf.set_text_color(*h2r('#374151'))
            pdf.multi_cell(RIGHT_W, 4, data['summary'])
            ry = pdf.get_y() + 4

        if data.get('experience'):
            ry = sec_title('EXPERIENCE', '#6366f1', ry, RIGHT_W, size=9, line_color='#6366f1')
            for e in data['experience']:
                pdf.set_xy(RIGHT_X, ry)
                pdf.set_font('Helvetica', 'B', 9.5)
                pdf.set_text_color(*h2r('#1f2937'))
                pdf.cell(RIGHT_W, 5, e.get('role',''), ln=1)
                pdf.set_xy(RIGHT_X, ry + 5)
                pdf.set_font('Helvetica', 'B', 8)
                pdf.set_text_color(*h2r('#6366f1'))
                pdf.cell(RIGHT_W, 4, e.get('company',''), ln=1)
                pdf.set_xy(RIGHT_X, ry + 9)
                pdf.set_fill_color(*h2r('#ede9fe'))
                pdf.set_text_color(*h2r('#6366f1'))
                pdf.set_font('Helvetica', 'B', 7)
                pdf.cell(0, 4, f"  {e.get('start','')}  -  {e.get('end','Present')}  ", ln=1, fill=True)
                ry += 14
                if e.get('description'):
                    pdf.set_xy(RIGHT_X + 3, ry)
                    pdf.set_font('Helvetica', '', 8)
                    pdf.set_text_color(*h2r('#374151'))
                    for line in e['description'].split('\n'):
                        line = line.strip().lstrip('* ').lstrip('-').strip()
                        if line:
                            pdf.multi_cell(RIGHT_W - 3, 3.5, f"*  {line}")
                    ry = pdf.get_y() + 2
                else:
                    ry += 2
            ry += 2

        if data.get('projects'):
            ry = sec_title('PROJECTS', '#6366f1', ry, RIGHT_W, size=9, line_color='#6366f1')
            for pr in data['projects']:
                pdf.set_xy(RIGHT_X, ry)
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(*h2r('#1f2937'))
                pdf.cell(RIGHT_W, 5, pr.get('name',''), ln=1)
                if pr.get('technologies'):
                    pdf.set_xy(RIGHT_X, ry + 5)
                    pdf.set_font('Helvetica', 'B', 7.5)
                    pdf.set_text_color(*h2r('#14b8a6'))
                    pdf.cell(RIGHT_W, 4, pr['technologies'], ln=1)
                    ry += 4
                if pr.get('description'):
                    pdf.set_xy(RIGHT_X + 3, ry + 5)
                    pdf.set_font('Helvetica', '', 8)
                    pdf.set_text_color(*h2r('#374151'))
                    pdf.multi_cell(RIGHT_W - 3, 3.5, pr['description'])
                    ry = pdf.get_y()
                ry += 8

    # ═══════════════════════════════════════════════════════════
    # EXECUTIVE
    # ═══════════════════════════════════════════════════════════
    elif template == 'executive':
        pdf.set_fill_color(*h2r('#0f4c3a'))
        pdf.rect(M, 10, CW, 28, style='F')
        pdf.set_xy(M, 16)
        pdf.set_font('Times', 'B', 20)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(CW, 8, name, ln=1, align='C')
        pdf.set_font('Times', '', 9)
        pdf.set_text_color(*h2r('#d4a017'))
        pdf.cell(CW, 5, '   |   '.join(contacts), ln=1, align='C')

        y = 44
        pdf.set_draw_color(*h2r('#d4a017'))
        pdf.set_line_width(0.6)
        pdf.line(M, y, PW - M, y)
        y += 6

        def exec_sec(text, y_pos):
            pdf.set_xy(M, y_pos)
            pdf.set_font('Times', 'B', 10)
            pdf.set_text_color(*h2r('#0f4c3a'))
            pdf.set_char_spacing(0.8)
            pdf.cell(CW, 6, text.upper(), ln=1)
            pdf.set_char_spacing(0)
            pdf.set_draw_color(*h2r('#e5e7eb'))
            pdf.set_line_width(0.2)
            pdf.line(M, y_pos + 6, PW - M, y_pos + 6)
            return y_pos + 9

        if data.get('summary'):
            y = exec_sec('Professional Summary', y)
            pdf.set_xy(M, y)
            pdf.set_font('Times', '', 9.5)
            pdf.set_text_color(*h2r('#374151'))
            pdf.multi_cell(CW, 5, data['summary'])
            y = pdf.get_y() + 4
            pdf.set_draw_color(*h2r('#d4a017'))
            pdf.set_line_width(0.6)
            pdf.line(M, y, PW - M, y)
            y += 6

        if data.get('experience'):
            y = exec_sec('Professional Experience', y)
            for e in data['experience']:
                pdf.set_xy(M + 4, y)
                pdf.set_draw_color(*h2r('#d4a017'))
                pdf.set_line_width(0.4)
                pdf.ellipse(M - 2, y + 1, 3, 3, style='F')
                pdf.set_font('Times', 'B', 10)
                pdf.set_text_color(*h2r('#1f2937'))
                pdf.cell(CW - 8, 5, e.get('role',''), ln=1)
                pdf.set_xy(M + 4, y + 5)
                pdf.set_font('Times', 'I', 9)
                pdf.set_text_color(*h2r('#0f4c3a'))
                pdf.cell(CW - 8, 4, e.get('company',''), ln=1)
                pdf.set_xy(M + 4, y + 9)
                pdf.set_font('Times', '', 8.5)
                pdf.set_text_color(*h2r('#6b7280'))
                pdf.cell(CW - 8, 4, f"{e.get('start','')}  -  {e.get('end','Present')}", ln=1)
                y += 14
                if e.get('description'):
                    pdf.set_xy(M + 8, y)
                    pdf.set_font('Times', '', 9)
                    pdf.set_text_color(*h2r('#374151'))
                    for line in e['description'].split('\n'):
                        line = line.strip().lstrip('* ').lstrip('-').strip()
                        if line:
                            pdf.multi_cell(CW - 12, 4, f" |  {line}")
                    y = pdf.get_y() + 2
                else:
                    y += 2
            y += 4
            pdf.set_draw_color(*h2r('#d4a017'))
            pdf.set_line_width(0.6)
            pdf.line(M, y, PW - M, y)
            y += 6

        if data.get('education'):
            y = exec_sec('Education', y)
            for e in data['education']:
                deg = f"{e.get('degree','')} in {e.get('field','')}" if e.get('field') else e.get('degree','')
                pdf.set_xy(M, y)
                pdf.set_font('Times', 'B', 10)
                pdf.set_text_color(*h2r('#1f2937'))
                pdf.cell(CW, 5, deg, ln=1)
                pdf.set_xy(M, y + 5)
                pdf.set_font('Times', 'I', 9)
                pdf.set_text_color(*h2r('#0f4c3a'))
                yr = f' {e["year"]}' if e.get('year') else ''
                pdf.cell(CW, 4, f"{e.get('institution','')}{yr}", ln=1)
                y += 12
            y += 4
            pdf.set_draw_color(*h2r('#d4a017'))
            pdf.set_line_width(0.6)
            pdf.line(M, y, PW - M, y)
            y += 6

        if data.get('skills'):
            y = exec_sec('Core Competencies', y)
            pdf.set_xy(M, y)
            pdf.set_font('Times', '', 9.5)
            pdf.set_text_color(*h2r('#374151'))
            for s in data['skills']:
                pdf.cell(pdf.get_string_width(s) + 10, 5, f" |  {s}  ", ln=0)
            pdf.ln()
            y = pdf.get_y() + 4
            pdf.set_draw_color(*h2r('#d4a017'))
            pdf.set_line_width(0.6)
            pdf.line(M, y, PW - M, y)
            y += 6

        if data.get('certifications'):
            y = exec_sec('Certifications', y)
            for c in data['certifications']:
                pdf.set_xy(M, y)
                pdf.set_font('Times', '', 9)
                pdf.set_text_color(*h2r('#374151'))
                org = f'  -  {c["org"]}' if c.get('org') else ''
                yr  = f' ({c["year"]})' if c.get('year') else ''
                pdf.cell(CW, 4, f"{c.get('name','')}{org}{yr}", ln=1)
                y += 5
            y += 4
            pdf.set_draw_color(*h2r('#d4a017'))
            pdf.set_line_width(0.6)
            pdf.line(M, y, PW - M, y)
            y += 6

        if data.get('languages'):
            y = exec_sec('Languages', y)
            pdf.set_xy(M, y)
            pdf.set_font('Times', '', 9.5)
            pdf.set_text_color(*h2r('#374151'))
            langs = [f"{l.get('language','')} ({l.get('level','')})" for l in data['languages']]
            pdf.cell(CW, 5, '   |   '.join(langs), ln=1)

    # ═══════════════════════════════════════════════════════════
    # CREATIVE
    # ═══════════════════════════════════════════════════════════
    elif template == 'creative':
        pdf.set_fill_color(*h2r('#8B3A3A'))
        pdf.rect(M, 10, LEFT_W, 277, style='F')
        ly, ry = 18, 18

        # Avatar initials
        initials = ''.join([w[0] for w in name.split() if w])[:2].upper() if name else '?'
        pdf.set_xy(LEFT_X + 10, ly)
        pdf.set_font('Helvetica', 'B', 14)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(LEFT_W - 20, 8, initials, ln=1, align='C')
        ly += 12

        if name:
            pdf.set_xy(LEFT_X + 3, ly)
            pdf.set_font('Helvetica', 'B', 10)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(LEFT_W - 6, 5, name, ln=1, align='C')
            ly += 6

        first_role = ''
        if data.get('experience'):
            first_role = next((e.get('role','') for e in data['experience'] if e.get('role')), '')
        if first_role:
            pdf.set_xy(LEFT_X + 3, ly)
            pdf.set_font('Helvetica', '', 7.5)
            pdf.set_text_color(200, 200, 210)
            pdf.cell(LEFT_W - 6, 4, first_role, ln=1, align='C')
            ly += 6

        if contacts:
            pdf.set_xy(LEFT_X + 3, ly)
            pdf.set_font('Helvetica', 'B', 7)
            pdf.set_text_color(180, 180, 190)
            pdf.cell(LEFT_W - 6, 5, 'CONTACT', ln=1)
            ly += 6
            pdf.set_font('Helvetica', '', 7)
            pdf.set_text_color(220, 220, 220)
            for c in contacts[:4]:
                pdf.set_xy(LEFT_X + 3, ly)
                pdf.cell(LEFT_W - 6, 4, c, ln=1)
                ly += 4.5
            ly += 2

        if data.get('skills'):
            pdf.set_xy(LEFT_X + 3, ly)
            pdf.set_font('Helvetica', 'B', 7)
            pdf.set_text_color(180, 180, 190)
            pdf.cell(LEFT_W - 6, 5, 'SKILLS', ln=1)
            ly += 6
            for i, s in enumerate(data['skills']):
                pdf.set_xy(LEFT_X + 3, ly)
                pdf.set_font('Helvetica', '', 7)
                pdf.set_text_color(255, 255, 255)
                pdf.cell(LEFT_W - 6, 4, s, ln=1)
                ly += 4
                bar_w = LEFT_W - 10
                fill_w = bar_w * max(95 - i * 7, 40) / 100
                pdf.set_fill_color(140, 80, 80)
                pdf.rect(LEFT_X + 3, ly, bar_w, 2, style='F')
                pdf.set_fill_color(220, 220, 220)
                pdf.rect(LEFT_X + 3, ly, fill_w, 2, style='F')
                ly += 4
            ly += 2

        if data.get('languages'):
            pdf.set_xy(LEFT_X + 3, ly)
            pdf.set_font('Helvetica', 'B', 7)
            pdf.set_text_color(180, 180, 190)
            pdf.cell(LEFT_W - 6, 5, 'LANGUAGES', ln=1)
            ly += 6
            for l in data['languages']:
                pdf.set_xy(LEFT_X + 3, ly)
                pdf.set_font('Helvetica', '', 7)
                pdf.set_text_color(220, 220, 220)
                pdf.cell(LEFT_W - 6, 4, f"{l.get('language','')}  {l.get('level','')}", ln=1)
                ly += 4.5
            ly += 2

        if data.get('certifications'):
            pdf.set_xy(LEFT_X + 3, ly)
            pdf.set_font('Helvetica', 'B', 7)
            pdf.set_text_color(180, 180, 190)
            pdf.cell(LEFT_W - 6, 5, 'CERTIFICATIONS', ln=1)
            ly += 6
            for c in data['certifications']:
                pdf.set_xy(LEFT_X + 3, ly)
                pdf.set_font('Helvetica', '', 7)
                pdf.set_text_color(220, 220, 220)
                yr = f' ({c["year"]})' if c.get('year') else ''
                pdf.cell(LEFT_W - 6, 4, f"{c.get('name','')}{yr}", ln=1)
                ly += 4.5
            ly += 2

        # ── Main ──
        if name:
            pdf.set_xy(RIGHT_X, ry)
            pdf.set_font('Helvetica', 'B', 16)
            pdf.set_text_color(*h2r('#8B3A3A'))
            pdf.cell(RIGHT_W, 7, name, ln=1)
            ry += 8
        if first_role:
            pdf.set_xy(RIGHT_X, ry)
            pdf.set_font('Helvetica', 'B', 9)
            pdf.set_text_color(*h2r('#8B3A3A'))
            pdf.cell(RIGHT_W, 5, first_role, ln=1)
            ry += 7

        if data.get('summary'):
            ry = sec_title('ABOUT ME', '#8B3A3A', ry, RIGHT_W, size=9, line_color='#f9e0e0')
            pdf.set_xy(RIGHT_X, ry)
            pdf.set_font('Helvetica', '', 8.5)
            pdf.set_text_color(*h2r('#374151'))
            pdf.multi_cell(RIGHT_W, 4, data['summary'])
            ry = pdf.get_y() + 4

        if data.get('experience'):
            ry = sec_title('EXPERIENCE', '#8B3A3A', ry, RIGHT_W, size=9, line_color='#f9e0e0')
            for e in data['experience']:
                pdf.set_xy(RIGHT_X, ry)
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(*h2r('#1f2937'))
                pdf.cell(RIGHT_W, 5, e.get('role',''), ln=1)
                pdf.set_xy(RIGHT_X, ry + 5)
                pdf.set_font('Helvetica', 'B', 8)
                pdf.set_text_color(*h2r('#8B3A3A'))
                pdf.cell(RIGHT_W, 4, e.get('company',''), ln=1)
                pdf.set_xy(RIGHT_X, ry + 9)
                pdf.set_font('Helvetica', '', 7.5)
                pdf.set_text_color(*h2r('#9ca3af'))
                pdf.cell(RIGHT_W, 4, f"{e.get('start','')}  -  {e.get('end','Present')}", ln=1)
                ry += 14
                if e.get('description'):
                    pdf.set_xy(RIGHT_X + 3, ry)
                    pdf.set_font('Helvetica', '', 8)
                    pdf.set_text_color(*h2r('#374151'))
                    for line in e['description'].split('\n'):
                        line = line.strip().lstrip('* ').lstrip('-').strip()
                        if line:
                            pdf.multi_cell(RIGHT_W - 3, 3.5, f">  {line}")
                    ry = pdf.get_y() + 2
                else:
                    ry += 2
            ry += 2

        if data.get('education'):
            ry = sec_title('EDUCATION', '#8B3A3A', ry, RIGHT_W, size=9, line_color='#f9e0e0')
            for e in data['education']:
                deg = f"{e.get('degree','')} in {e.get('field','')}" if e.get('field') else e.get('degree','')
                pdf.set_xy(RIGHT_X, ry)
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(*h2r('#1f2937'))
                pdf.cell(RIGHT_W, 5, deg, ln=1)
                pdf.set_xy(RIGHT_X, ry + 5)
                pdf.set_font('Helvetica', '', 8)
                pdf.set_text_color(*h2r('#6b7280'))
                yr = f' {e["year"]}' if e.get('year') else ''
                pdf.cell(RIGHT_W, 4, f"{e.get('institution','')}{yr}", ln=1)
                ry += 11
            ry += 2

        if data.get('projects'):
            ry = sec_title('PROJECTS', '#8B3A3A', ry, RIGHT_W, size=9, line_color='#f9e0e0')
            for pr in data['projects']:
                pdf.set_xy(RIGHT_X, ry)
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(*h2r('#1f2937'))
                pdf.cell(RIGHT_W, 5, pr.get('name',''), ln=1)
                if pr.get('technologies'):
                    pdf.set_xy(RIGHT_X, ry + 5)
                    pdf.set_font('Helvetica', 'B', 7.5)
                    pdf.set_text_color(*h2r('#8B3A3A'))
                    pdf.cell(RIGHT_W, 4, pr['technologies'], ln=1)
                    ry += 4
                if pr.get('description'):
                    pdf.set_xy(RIGHT_X + 3, ry + 5)
                    pdf.set_font('Helvetica', '', 8)
                    pdf.set_text_color(*h2r('#374151'))
                    pdf.multi_cell(RIGHT_W - 3, 3.5, pr['description'])
                    ry = pdf.get_y()
                ry += 8

    # ═══════════════════════════════════════════════════════════
    # PROFESSIONAL (default)
    # ═══════════════════════════════════════════════════════════
    else:
        pdf.set_xy(M, 15)
        pdf.set_font('Helvetica', 'B', 18)
        pdf.set_text_color(*h2r('#1e1b4b'))
        pdf.cell(CW, 8, name, ln=1, align='C')
        pdf.set_font('Helvetica', '', 8)
        pdf.set_text_color(*h2r('#6b7280'))
        pdf.cell(CW, 5, '  |  '.join(contacts), ln=1, align='C')
        pdf.set_draw_color(*h2r('#a78bfa'))
        pdf.set_line_width(0.8)
        pdf.line(M, 32, PW - M, 32)
        y = 38

        if data.get('summary'):
            y = sec_title('PROFESSIONAL SUMMARY', '#7c3aed', y, CW, size=10, line_color='#ede9fe')
            pdf.set_xy(M, y)
            pdf.set_font('Helvetica', '', 9)
            pdf.set_text_color(*h2r('#374151'))
            pdf.multi_cell(CW, 4.5, data['summary'])
            y = pdf.get_y() + 4

        if data.get('experience'):
            y = sec_title('WORK EXPERIENCE', '#7c3aed', y, CW, size=10, line_color='#ede9fe')
            for e in data['experience']:
                y = entry_block(
                    f"{e.get('role','')}{'  -  ' if e.get('role') and e.get('company') else ''}{e.get('company','')}",
                    f"{e.get('start','')}  -  {e.get('end','Present')}",
                    e.get('description',''), y, CW,
                    title_color='#1e1b4b', sub_color='#6b7280', body_color='#374151'
                )
            y += 2

        # Two-column: Education + Projects | Skills + Certs + Langs
        mid = y
        if data.get('education') or data.get('projects'):
            y = sec_title('EDUCATION', '#7c3aed', y, CW / 2 - 4, size=10, line_color='#ede9fe')
            for e in data['education']:
                deg = f"{e.get('degree','')} in {e.get('field','')}" if e.get('field') else e.get('degree','')
                pdf.set_xy(M, y)
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(*h2r('#1e1b4b'))
                pdf.cell(CW / 2 - 4, 5, deg, ln=1)
                pdf.set_xy(M, y + 5)
                pdf.set_font('Helvetica', 'I', 8)
                pdf.set_text_color(*h2r('#6b7280'))
                yr = f' {e["year"]}' if e.get('year') else ''
                pdf.cell(CW / 2 - 4, 4, f"{e.get('institution','')}{yr}", ln=1)
                y += 11
            if data.get('projects'):
                y = sec_title('PROJECTS', '#7c3aed', y, CW / 2 - 4, size=10, line_color='#ede9fe')
                for pr in data['projects']:
                    pdf.set_xy(M, y)
                    pdf.set_font('Helvetica', 'B', 9)
                    pdf.set_text_color(*h2r('#1e1b4b'))
                    pdf.cell(CW / 2 - 4, 5, pr.get('name',''), ln=1)
                    if pr.get('technologies'):
                        pdf.set_xy(M, y + 5)
                        pdf.set_font('Helvetica', 'I', 8)
                        pdf.set_text_color(*h2r('#6b7280'))
                        pdf.cell(CW / 2 - 4, 4, f"Tech: {pr['technologies']}", ln=1)
                        y += 5
                    if pr.get('description'):
                        pdf.set_xy(M, y + 5)
                        pdf.set_font('Helvetica', '', 8)
                        pdf.set_text_color(*h2r('#374151'))
                        pdf.multi_cell(CW / 2 - 4, 3.5, pr['description'])
                        y = pdf.get_y()
                    y += 8
            y += 4

        ry = mid
        if data.get('skills'):
            ry = sec_title('SKILLS', '#7c3aed', ry, CW / 2 - 4, size=10, line_color='#ede9fe')
            pdf.set_xy(M + CW / 2 + 4, ry)
            pdf.set_font('Helvetica', '', 8.5)
            pdf.set_text_color(*h2r('#374151'))
            for s in data['skills']:
                pdf.set_fill_color(*h2r('#f5f3ff'))
                pdf.set_draw_color(*h2r('#ede9fe'))
                tw = pdf.get_string_width(s) + 10
                pdf.cell(tw, 5, f'  {s}  ', ln=0, fill=True, border=1)
                pdf.set_xy(pdf.get_x() + 2, ry)
            pdf.ln()
            ry = pdf.get_y() + 4

        if data.get('certifications'):
            ry = sec_title('CERTIFICATIONS', '#7c3aed', ry, CW / 2 - 4, size=10, line_color='#ede9fe')
            for c in data['certifications']:
                pdf.set_xy(M + CW / 2 + 4, ry)
                pdf.set_font('Helvetica', '', 8.5)
                pdf.set_text_color(*h2r('#374151'))
                org = f'  -  {c["org"]}' if c.get('org') else ''
                yr  = f' ({c["year"]})' if c.get('year') else ''
                pdf.cell(CW / 2 - 4, 4, f"{c.get('name','')}{org}{yr}", ln=1)
                ry += 5
            ry += 2

        if data.get('languages'):
            ry = sec_title('LANGUAGES', '#7c3aed', ry, CW / 2 - 4, size=10, line_color='#ede9fe')
            pdf.set_xy(M + CW / 2 + 4, ry)
            pdf.set_font('Helvetica', '', 8.5)
            pdf.set_text_color(*h2r('#374151'))
            langs = [f"{l.get('language','')} ({l.get('level','')})" for l in data['languages']]
            pdf.cell(CW / 2 - 4, 5, '   |   '.join(langs), ln=1)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def _generate_docx(data):
    """Generate a clean ATS-friendly DOCX using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    purple = RGBColor(0xa7, 0x8b, 0xfa)
    dark   = RGBColor(0x1e, 0x1b, 0x4b)
    gray   = RGBColor(0x6b, 0x72, 0x80)

    def add_name(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = dark
        p.paragraph_format.space_after = Pt(4)

    def add_contact(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = gray
        p.paragraph_format.space_after = Pt(6)

    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = purple
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        # underline via border - add bottom border
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'a78bfa')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_body(text, bold=False, italic=False, size=9.5, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color or dark
        p.paragraph_format.space_after = Pt(2)
        return p

    p = data.get('personal', {})
    if p.get('name'):
        add_name(p['name'])

    contacts = [x for x in [p.get('email'), p.get('phone'), p.get('location'), p.get('linkedin'), p.get('portfolio')] if x]
    if contacts:
        add_contact(' | '.join(contacts))

    if data.get('summary'):
        add_section_heading('PROFESSIONAL SUMMARY')
        add_body(data['summary'])

    if data.get('experience'):
        add_section_heading('WORK EXPERIENCE')
        for e in data['experience']:
            if e.get('role') or e.get('company'):
                add_body(f"{e.get('role','')} — {e.get('company','')}", bold=True)
                add_body(f"{e.get('start','')} – {e.get('end','Present')}", italic=True, color=gray)
                if e.get('description'):
                    for line in e['description'].split('\n'):
                        line = line.strip().lstrip('•-').strip()
                        if line:
                            add_body(f'• {line}')

    if data.get('education'):
        add_section_heading('EDUCATION')
        for e in data['education']:
            if e.get('institution') or e.get('degree'):
                deg = f"{e.get('degree','')} in {e.get('field','')}" if e.get('field') else e.get('degree','')
                add_body(deg, bold=True)
                add_body(f"{e.get('institution','')} {('· ' + str(e['year'])) if e.get('year') else ''}", color=gray)

    if data.get('skills'):
        add_section_heading('SKILLS')
        add_body(', '.join(data['skills']))

    if data.get('projects'):
        add_section_heading('PROJECTS')
        for pr in data['projects']:
            if pr.get('name'):
                add_body(pr['name'], bold=True)
                if pr.get('technologies'):
                    add_body(f"Technologies: {pr['technologies']}", italic=True, color=gray)
                if pr.get('description'):
                    add_body(pr['description'])

    if data.get('certifications'):
        add_section_heading('CERTIFICATIONS')
        for c in data['certifications']:
            if c.get('name'):
                add_body(f"{c['name']} — {c.get('org','')} {c.get('year','')}")

    if data.get('languages'):
        add_section_heading('LANGUAGES')
        langs = [f"{l.get('language','')} ({l.get('level','')})" for l in data['languages'] if l.get('language')]
        add_body(', '.join(langs))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
